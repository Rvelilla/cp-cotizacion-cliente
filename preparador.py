import os
from docx import Document

def preparar_plantilla():
    print("Iniciando preparación de la plantilla...")
    
    archivo_origen = None
    for archivo in os.listdir('.'):
        if archivo.endswith('.docx') and not archivo.startswith('~$') and ('Cliente' in archivo or 'Furgon' in archivo or 'base' in archivo):
            archivo_origen = archivo
            break
            
    if not archivo_origen:
        print("ERROR: No encontré el archivo base.")
        return False
        
    print(f"✔️ Archivo detectado exitosamente: '{archivo_origen}'")
    
    doc = Document(archivo_origen)
    
    # 1. Reemplazar encabezado y textos
    for para in doc.paragraphs:
        if 'Medellín' in para.text or 'mayo 29' in para.text:
            para.text = "{{ ciudad }}, {{ fecha }}"
        if 'FAKE' in para.text:
            para.text = para.text.replace('FAKE', '{{ nombre_cliente }}')
            
        if 'Garantía del la estaca es de un año' in para.text or 'Garantía de la estaca' in para.text:
            para.text = para.text.replace('Garantía del la estaca es de un año.', '{{ texto_garantia }}')
            para.text = para.text.replace('Garantía de la estaca es de un año.', '{{ texto_garantia }}')

        if 'Claudia Velilla' in para.text:
            para.text = para.text.replace('Claudia Velilla', '{{ nombre_asesor }}')
        if 'Dir. Mercadeo' in para.text:
            para.text = para.text.replace('Dir. Mercadeo', '{{ cargo_asesor }}')

    # 2. Reemplazar Tabla de Medidas (Corregido para 8 columnas)
    if len(doc.tables) > 0:
        tabla_medidas = doc.tables[0]
        if len(tabla_medidas.rows) > 2:
            celdas = tabla_medidas.rows[2].cells
            # Ahora asignamos cada variable a su columna exacta
            if len(celdas) >= 8:
                celdas[0].text = '{{largo_ext}}'
                celdas[1].text = '{{ancho_ext}}'
                celdas[2].text = '{{alto_ext}}'
                celdas[3].text = '{{tipo_carroceria}}'
                celdas[4].text = '{{peso_kg}} (+/-36)'
                celdas[5].text = '{{valor_c_rra}}'
                celdas[6].text = '{{iva}}'
                celdas[7].text = '{{subtotal}}'

    # 3. Reemplazar Tabla de Especificaciones
    if len(doc.tables) > 1:
        tabla_esp = doc.tables[1]
        while len(tabla_esp.rows) > 2:
            tr = tabla_esp.rows[2]._tr
            tr.getparent().remove(tr)
            
        if len(tabla_esp.rows) > 1:
            celdas_esp = tabla_esp.rows[1].cells
            celdas_esp[0].text = '{% for item in especificaciones %}{{item.concepto}}'
            celdas_esp[1].text = '{{item.detalle}}{% endfor %}'

    os.makedirs('plantillas', exist_ok=True)
    doc.save("plantillas/plantilla_cotizacion.docx")
    print("¡Éxito! Plantilla finalizada con todas las variables.")
    return True

if __name__ == '__main__':
    preparar_plantilla()